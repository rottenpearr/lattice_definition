import docx
from docx.shared import Inches
from datetime import datetime


def save_docx(text, img_path):
    doc = docx.Document()

    doc.add_heading("Результаты поиска кристаллической решетки", level=1)

    date = datetime.now().strftime("%d %B %Y %H-%M-%S")
    doc.add_paragraph(f"Дата: {date}")

    doc.add_picture(str(img_path), width=Inches(2))

    for row in text.split('\n'):
        doc.add_paragraph(row)

    doc.save(f"reports/Отчет {date}.docx")

    print(f"Документ сохранен как 'Отчет {date}.docx' в папку 'reports'.")
