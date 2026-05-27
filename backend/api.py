"""
CRIS — FastAPI backend
Запуск: uvicorn backend.api:app --reload --port 8002
Документация: http://localhost:8002/docs
"""

import hashlib
import json
import os
import sys

from pathlib import Path
from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env")

import io
from datetime import datetime

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional

from cris.core.coordinates import shift_coordinates, normalize_coordinates
from cris.db.connection import get_cursor
from cris.db.queries import check_coords

# ── GigaChat (опционально) ────────────────────────────────────────────────────
try:
    from gigachat import GigaChat
    from gigachat.models import Chat, Messages, MessagesRole
    _GIGACHAT_AVAILABLE = True
except ImportError:
    _GIGACHAT_AVAILABLE = False

_GC_CREDENTIALS = os.getenv("GIGACHAT_CREDENTIALS", "")
_GC_SCOPE       = os.getenv("GIGACHAT_SCOPE", "GIGACHAT_API_PERS")
_GC_MODEL       = os.getenv("GIGACHAT_MODEL", "GigaChat-2-Max")
_GC_READY       = _GIGACHAT_AVAILABLE and bool(_GC_CREDENTIALS)

# ── Приложение ────────────────────────────────────────────────────────────────

app = FastAPI(
    title="CRIS API",
    description="Crystal Recognition & Identification System",
    version="0.4.3",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Схемы ─────────────────────────────────────────────────────────────────────

class Ion(BaseModel):
    label: str
    x: float
    y: float
    z: float


class AnalyzeRequest(BaseModel):
    sites: list[Ion]
    methods: list[str] = ["db"]   # "db" | "catboost" | "catboost_substance" | "rf"


class LatticeResult(BaseModel):
    id: Optional[int]
    name_en: Optional[str]
    name_ru: Optional[str]
    description: Optional[str]
    confidence: float


class StructureResult(BaseModel):
    id: Optional[int]
    name: Optional[str]
    formula: Optional[str]
    cell_a: Optional[float]
    cell_b: Optional[float]
    cell_c: Optional[float]
    sg_number: Optional[int]
    sg_hm: Optional[str]
    confidence: float


class RankingItem(BaseModel):
    name_en: Optional[str]
    name_ru: Optional[str]
    prob: float


class MLMethodResult(BaseModel):
    method: str                      # "catboost" | "catboost_substance" | "rf"
    category: str = "crystal_system" # "crystal_system" | "substance"
    name_en: Optional[str] = None
    name_ru: Optional[str] = None
    confidence: float = 0.0
    ranking: list[RankingItem] = []


class AnalyzeResponse(BaseModel):
    success: bool
    message: str = ""
    session_id: Optional[str] = None
    lattice: Optional[LatticeResult] = None
    structure: Optional[StructureResult] = None
    lattice_ranking: list[RankingItem] = []
    ml_results: list[MLMethodResult] = []


class ChatMessage(BaseModel):
    role: str       # "user" | "assistant"
    content: str


class ChatRequest(BaseModel):
    session_id: str                 # UUID из /api/analyze
    messages: list[ChatMessage]     # вся история диалога (фронт хранит локально)


class ChatResponse(BaseModel):
    reply: str
    model: str


# ── Вспомогательные функции ───────────────────────────────────────────────────

import re as _re

_SUBSCRIPT_MAP   = str.maketrans('0123456789+-', '₀₁₂₃₄₅₆₇₈₉₊₋')
_SUPERSCRIPT_MAP = str.maketrans('0123456789+-', '⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻')

# LaTeX команды → Unicode символы
_LATEX_COMMANDS = [
    # Греческие буквы
    (r'\\alpha',   'α'), (r'\\beta',    'β'), (r'\\gamma',   'γ'),
    (r'\\delta',   'δ'), (r'\\epsilon', 'ε'), (r'\\zeta',    'ζ'),
    (r'\\eta',     'η'), (r'\\theta',   'θ'), (r'\\iota',    'ι'),
    (r'\\kappa',   'κ'), (r'\\lambda',  'λ'), (r'\\mu',      'μ'),
    (r'\\nu',      'ν'), (r'\\xi',      'ξ'), (r'\\pi',      'π'),
    (r'\\rho',     'ρ'), (r'\\sigma',   'σ'), (r'\\tau',     'τ'),
    (r'\\upsilon', 'υ'), (r'\\phi',     'φ'), (r'\\chi',     'χ'),
    (r'\\psi',     'ψ'), (r'\\omega',   'ω'),
    (r'\\Gamma',   'Γ'), (r'\\Delta',   'Δ'), (r'\\Theta',   'Θ'),
    (r'\\Lambda',  'Λ'), (r'\\Xi',      'Ξ'), (r'\\Pi',      'Π'),
    (r'\\Sigma',   'Σ'), (r'\\Phi',     'Φ'), (r'\\Psi',     'Ψ'),
    (r'\\Omega',   'Ω'),
    # Единицы и спецсимволы
    (r'\\AA\b',    'Å'), (r'\\angstrom','Å'),
    (r'\\degree',  '°'), (r'\\circ',    '°'),
    (r'\\cdot',    '·'), (r'\\times',   '×'),
    (r'\\pm',      '±'), (r'\\mp',      '∓'),
    (r'\\leq',     '≤'), (r'\\geq',     '≥'),
    (r'\\neq',     '≠'), (r'\\approx',  '≈'),
    (r'\\infty',   '∞'), (r'\\partial', '∂'),
    (r'\\nabla',   '∇'), (r'\\sqrt',    '√'),
    # Пробелы LaTeX
    (r'\\,',  ''), (r'\\;',  ' '), (r'\\ ',  ' '), (r'\\!',  ''),
    (r'\\quad', ' '), (r'\\qquad', '  '),
]


def _latex_to_text(expr: str) -> str:
    """Конвертирует LaTeX-выражение в читаемый Unicode-текст."""
    s = expr

    # \text{...} \mathrm{...} \mathbf{...} и т.п. → содержимое
    s = _re.sub(r'\\(?:text|mathrm|mathbf|mathit|mathsf|mbox|operatorname)\{([^}]*)\}', r'\1', s)

    # {,} → ,  и  {.} → .  (европейский разделитель в числах)
    s = _re.sub(r'\{([.,])\}', r'\1', s)

    # Убираем лишние фигурные скобки вокруг одного символа: {a} → a
    s = _re.sub(r'\{([^{}])\}', r'\1', s)

    # Команды → Unicode
    for pattern, repl in _LATEX_COMMANDS:
        s = _re.sub(pattern, repl, s)

    # Подстрочные: _{...} или _X
    def _do_sub(m):
        inner = (m.group(1) or m.group(2) or '').strip('{}')
        return inner.translate(_SUBSCRIPT_MAP)
    s = _re.sub(r'_\{([^}]*)\}|_([^\s{\\])', _do_sub, s)

    # Надстрочные: ^{...} или ^X
    def _do_sup(m):
        inner = (m.group(1) or m.group(2) or '').strip('{}')
        return inner.translate(_SUPERSCRIPT_MAP)
    s = _re.sub(r'\^\{([^}]*)\}|\^([^\s{\\])', _do_sup, s)

    # Оставшиеся фигурные скобки убираем
    s = s.replace('{', '').replace('}', '')

    # Оставшиеся \команды без аргументов убираем
    s = _re.sub(r'\\[a-zA-Z]+\*?', '', s)

    return s.strip()


def _fix_gigachat_unicode(text: str) -> str:
    """Заменяет GigaChat-специфичные escape-последовательности и LaTeX-паттерны на Unicode."""
    # 1. \unicodexXXXX → chr()
    def _uni_sub(m):
        try:
            return chr(int(m.group(1), 16))
        except ValueError:
            return m.group(0)
    text = _re.sub(r'\\unicodex([0-9A-Fa-f]{4,6})', _uni_sub, text)

    # 2. Блоки $$...$$ (display math)
    text = _re.sub(r'\$\$(.+?)\$\$', lambda m: _latex_to_text(m.group(1)), text, flags=_re.DOTALL)

    # 3. Инлайн $...$
    text = _re.sub(r'\$([^$\n]+?)\$', lambda m: _latex_to_text(m.group(1)), text)

    # 4. Одиночные LaTeX-команды вне $$ (GigaChat иногда пишет их без долларов)
    text = _latex_to_text(text)

    return text


def _input_hash(normalized: list) -> str:
    """SHA-256 нормализованных координат — стабильный идентификатор входа."""
    coords = sorted([[round(r[1], 6), round(r[2], 6), round(r[3], 6)] for r in normalized])
    return hashlib.sha256(json.dumps(coords).encode()).hexdigest()


def _find_or_create_session(input_hash: str, ion_count: int) -> str:
    """
    Ищет существующую сессию по хешу координат.
    Если не найдена — создаёт новую.
    Возвращает session_id (UUID строкой).
    """
    with get_cursor() as cur:
        cur.execute(
            "SELECT id FROM recognition_session WHERE input_hash = %s ORDER BY created_at DESC LIMIT 1",
            (input_hash,)
        )
        row = cur.fetchone()
        if row:
            return str(row[0])
        cur.execute(
            "INSERT INTO recognition_session (ion_count, input_hash) VALUES (%s, %s) RETURNING id",
            (ion_count, input_hash)
        )
        return str(cur.fetchone()[0])


def _save_recognition_result(session_id: str, lattice_type_id: Optional[int],
                              structure_id: Optional[int], confidence: float,
                              method: str = "GEOMETRIC"):
    """Сохраняет результат анализа в recognition_result."""
    with get_cursor() as cur:
        cur.execute(
            """
            INSERT INTO recognition_result
                (session_id, method, method_version, rank,
                 predicted_lattice_type_id, predicted_structure_id, confidence)
            VALUES (%s, %s, '1.0', 1, %s, %s, %s)
            ON CONFLICT (session_id, method, method_version, rank)
            DO UPDATE SET
                predicted_lattice_type_id = EXCLUDED.predicted_lattice_type_id,
                predicted_structure_id    = EXCLUDED.predicted_structure_id,
                confidence                = EXCLUDED.confidence,
                computed_at               = NOW()
            """,
            (session_id, method, lattice_type_id, structure_id, confidence)
        )


def _load_session_context(session_id: str) -> dict:
    """
    Загружает богатый контекст для системного промпта GigaChat.
    JOIN: recognition_result → lattice_type + lattice_metadata → reference_structure + substance_info
    """
    with get_cursor() as cur:
        cur.execute(
            """
            SELECT
                lt.name_ru,
                lt.name_en,
                lt.crystal_system,
                lt.bravais_lattice,
                lt.description            AS lt_description,
                lm.coordination_number,
                lm.packing_efficiency,
                lm.typical_materials,
                lm.applications           AS lt_applications,
                rs.name                   AS struct_name,
                rs.formula,
                rs.sg_hm,
                rs.sg_number,
                rs.cell_length_a,
                rs.cell_length_b,
                rs.cell_length_c,
                rs.cod_id,
                rs.mp_id,
                rs.doi,
                si.description            AS substance_desc,
                si.applications           AS substance_apps,
                si.hazards,
                si.properties,
                rr.confidence,
                rs_session.ion_count
            FROM recognition_result rr
            JOIN recognition_session rs_session ON rs_session.id = rr.session_id
            LEFT JOIN lattice_type lt       ON lt.id  = rr.predicted_lattice_type_id
            LEFT JOIN lattice_metadata lm   ON lm.lattice_type_id = lt.id
            LEFT JOIN reference_structure rs ON rs.id = rr.predicted_structure_id
            LEFT JOIN substance_info si     ON si.structure_id = rs.id
            WHERE rr.session_id = %s AND rr.method = 'GEOMETRIC' AND rr.rank = 1
            LIMIT 1
            """,
            (session_id,)
        )
        row = cur.fetchone()
        if not row:
            return {}
        cols = [d[0] for d in cur.description]
        return dict(zip(cols, row))


def _run_ml_methods(normalized: list, methods: list[str], session_id: Optional[str]) -> list[MLMethodResult]:
    """
    Запускает выбранные ML-методы и возвращает список MLMethodResult.
    Каждый результат также сохраняется в recognition_result.

    Категории:
      crystal_system — CatBoost (сингония): "catboost"
      substance      — CatBoost (вещество): "catboost_substance"; RF: "rf"
    """
    from cris.core.ml_predict import (
        predict_catboost, predict_catboost_substance, predict_rf, resolve_lattice_ids
    )

    results: list[MLMethodResult] = []

    def _preds_to_result(method: str, category: str, preds: list[dict]) -> MLMethodResult:
        if not preds:
            return MLMethodResult(method=method, category=category)
        top = preds[0]
        name_ru = None
        if top.get("lattice_type_id"):
            try:
                with get_cursor() as cur:
                    cur.execute("SELECT name_ru FROM lattice_type WHERE id=%s", (top["lattice_type_id"],))
                    row = cur.fetchone()
                    if row:
                        name_ru = row[0]
            except Exception:
                pass
        ranking = [
            RankingItem(name_en=p["lattice_name"], name_ru=None, prob=round(p["confidence"] * 100, 1))
            for p in preds
        ]
        return MLMethodResult(
            method=method,
            category=category,
            name_en=top["lattice_name"],
            name_ru=name_ru,
            confidence=round(top["confidence"] * 100, 2),
            ranking=ranking,
        )

    if "catboost" in methods:
        try:
            preds = predict_catboost(normalized)
            preds = resolve_lattice_ids(preds)
            mr = _preds_to_result("catboost", "crystal_system", preds)
            results.append(mr)
            if session_id and preds:
                _save_recognition_result(session_id, preds[0].get("lattice_type_id"), None,
                                         mr.confidence, method="CATBOOST")
        except Exception as e:
            logger.warning("CatBoost inference failed: {}", e)
            results.append(MLMethodResult(method="catboost", category="crystal_system"))

    if "catboost_substance" in methods:
        try:
            preds = predict_catboost_substance(normalized)
            preds = resolve_lattice_ids(preds)
            mr = _preds_to_result("catboost_substance", "substance", preds)
            results.append(mr)
            if session_id and preds:
                _save_recognition_result(session_id, preds[0].get("lattice_type_id"), None,
                                         mr.confidence, method="CATBOOST_SUBSTANCE")
        except Exception as e:
            logger.warning("CatBoost-substance inference failed: {}", e)
            results.append(MLMethodResult(method="catboost_substance", category="substance"))

    if "rf" in methods:
        try:
            preds = predict_rf(normalized)
            preds = resolve_lattice_ids(preds)
            mr = _preds_to_result("rf", "substance", preds)
            results.append(mr)
            if session_id and preds:
                _save_recognition_result(session_id, preds[0].get("lattice_type_id"), None,
                                         mr.confidence, method="RF")
        except Exception as e:
            logger.warning("RF inference failed: {}", e)
            results.append(MLMethodResult(method="rf", category="substance"))

    return results


def _build_system_prompt(ctx: dict) -> str:
    """Строит системный промпт из данных БД."""
    lines = [
        "Ты — ИИ-ассистент системы CRIS (Crystal Recognition & Identification System).",
        "Ты помогаешь учёным и инженерам разбираться в результатах распознавания кристаллических решёток.",
        "Отвечай научно, кратко и по существу. Используй русский язык.",
        "Пиши формулы и символы обычным текстом Unicode (Fe₃O₄, P4₂/mnm, 5.64 Å) — не используй LaTeX ($...$, \\text{}, \\, и подобное).",
    ]

    if not ctx:
        return "\n".join(lines)

    lines.append("\n═══ ТЕКУЩАЯ СТРУКТУРА ═══")

    # Тип решётки
    lt_name = ctx.get("name_ru") or ctx.get("name_en")
    if lt_name:
        en = ctx.get("name_en", "")
        lines.append(f"Тип решётки: {lt_name}" + (f" ({en})" if en and en != lt_name else ""))
    if ctx.get("crystal_system"):
        lines.append(f"Кристаллическая система: {ctx['crystal_system']}")
    if ctx.get("bravais_lattice"):
        lines.append(f"Решётка Браве: {ctx['bravais_lattice']}")
    if ctx.get("confidence") is not None:
        lines.append(f"Уверенность распознавания: {ctx['confidence']:.2f}")
    if ctx.get("ion_count"):
        lines.append(f"Число ионов в ячейке: {ctx['ion_count']}")

    # Метаданные типа решётки
    if ctx.get("coordination_number"):
        lines.append(f"Координационное число: {ctx['coordination_number']}")
    if ctx.get("packing_efficiency"):
        lines.append(f"Плотность упаковки: {ctx['packing_efficiency']:.2f}")
    if ctx.get("typical_materials"):
        lines.append(f"Типичные материалы: {ctx['typical_materials']}")
    if ctx.get("lt_applications"):
        lines.append(f"Применения типа решётки: {ctx['lt_applications']}")
    if ctx.get("lt_description"):
        lines.append(f"Описание типа: {ctx['lt_description']}")

    # Эталонная структура
    if ctx.get("struct_name") or ctx.get("formula"):
        lines.append("")
        lines.append("═══ ЭТАЛОННАЯ СТРУКТУРА ═══")
    if ctx.get("struct_name"):
        lines.append(f"Название: {ctx['struct_name']}")
    if ctx.get("formula"):
        lines.append(f"Формула: {ctx['formula']}")
    if ctx.get("sg_hm"):
        sg_n = ctx.get("sg_number", "")
        lines.append(f"Пространственная группа: {ctx['sg_hm']}" + (f" (№{sg_n})" if sg_n else ""))
    a, b, c = ctx.get("cell_length_a"), ctx.get("cell_length_b"), ctx.get("cell_length_c")
    if a and b and c:
        lines.append(f"Параметры ячейки: a={a:.3f} Å, b={b:.3f} Å, c={c:.3f} Å")
    if ctx.get("cod_id"):
        lines.append(f"COD ID: {ctx['cod_id']}")
    if ctx.get("mp_id"):
        lines.append(f"Materials Project: {ctx['mp_id']}")
    if ctx.get("doi"):
        lines.append(f"DOI: {ctx['doi']}")

    # Описание вещества из substance_info
    if ctx.get("substance_desc"):
        lines.append("")
        lines.append("═══ ВЕЩЕСТВО ═══")
        lines.append(ctx["substance_desc"])
    if ctx.get("substance_apps"):
        lines.append(f"Применение: {ctx['substance_apps']}")
    if ctx.get("hazards"):
        lines.append(f"Опасность: {ctx['hazards']}")
    if ctx.get("properties"):
        try:
            props = ctx["properties"]
            if isinstance(props, str):
                props = json.loads(props)
            if isinstance(props, dict) and props:
                lines.append("Свойства: " + "; ".join(f"{k}={v}" for k, v in list(props.items())[:6]))
        except Exception:
            pass

    lines.append(
        "\nОтвечай на вопросы пользователя, опираясь на данные выше."
        " Если в базе нет конкретных сведений по веществу или применению — используй свои научные знания,"
        " явно указывая, что это общая информация, а не данные из базы системы CRIS."
    )
    return "\n".join(lines)


def _save_chat_messages(session_id: str, user_content: str, assistant_content: str):
    """Сохраняет пару user/assistant в таблицу chat_message."""
    try:
        with get_cursor() as cur:
            cur.execute(
                "INSERT INTO chat_message (session_id, role, content) VALUES (%s, 'user', %s), (%s, 'assistant', %s)",
                (session_id, user_content, session_id, assistant_content)
            )
    except Exception:
        pass  # не ломаем ответ из-за ошибки логирования


# ── Эндпоинты ─────────────────────────────────────────────────────────────────

@app.get("/api/health")
def health():
    return {"status": "ok", "version": "0.4.3"}


@app.get("/api/debug/env")
def debug_env():
    return {
        "gigachat_available": _GIGACHAT_AVAILABLE,
        "gigachat_ready":     _GC_READY,
        "credentials_set":    bool(_GC_CREDENTIALS),
        "credentials_len":    len(_GC_CREDENTIALS),
        "scope":              _GC_SCOPE,
        "model":              _GC_MODEL,
        "env_file":           str(Path(__file__).parent.parent / ".env"),
        "env_file_exists":    (Path(__file__).parent.parent / ".env").exists(),
        "python_executable":  sys.executable,
        "python_version":     sys.version,
    }


@app.get("/api/stats")
def stats():
    try:
        with get_cursor() as cur:
            cur.execute("SELECT COUNT(*) FROM reference_structure")
            struct_count = cur.fetchone()[0]
            cur.execute("SELECT COUNT(*) FROM lattice_type")
            lattice_count = cur.fetchone()[0]
            cur.execute("SELECT COUNT(*) FROM recognition_session")
            session_count = cur.fetchone()[0]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DB error: {e}")

    return {
        "struct_count":   struct_count,
        "lattice_count":  lattice_count,
        "session_count":  session_count,
    }


@app.post("/api/analyze", response_model=AnalyzeResponse)
def analyze(body: AnalyzeRequest):
    """
    Принимает декартовые координаты ионов, запускает анализ.
    Создаёт/находит сессию, сохраняет результат в recognition_result.
    Возвращает session_id для последующего чата.
    """
    if len(body.sites) < 2:
        raise HTTPException(status_code=400, detail="Нужно минимум 2 иона")

    raw = [[s.label, s.x, s.y, s.z] for s in body.sites]

    try:
        shifted    = shift_coordinates(raw)
        normalized = normalize_coordinates(shifted)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))

    input_hash  = _input_hash(normalized)
    coords_dict = {i: row for i, row in enumerate(normalized)}
    result      = check_coords(coords_dict, len(body.sites))

    # ── Создаём / находим сессию ──────────────────────────────
    try:
        session_id = _find_or_create_session(input_hash, len(body.sites))
    except Exception:
        session_id = None

    # ── ML методы запускаем всегда, независимо от GEOMETRIC ──
    ml_results = _run_ml_methods(normalized, body.methods, session_id)

    if result is False:
        return AnalyzeResponse(
            success    = False,
            message    = "Совпадающих структур не найдено в эталонной БД",
            session_id = session_id,
            ml_results = ml_results,
        )

    lattice_names, struct_names = result[0]
    lattice_info, lt_prob       = result[1]
    struct_info,  st_prob       = result[2]

    lattice = LatticeResult(
        id          = lattice_info[0] if lattice_info else None,
        name_en     = lattice_info[1] if lattice_info else None,
        name_ru     = lattice_info[2] if lattice_info else None,
        description = lattice_info[3] if lattice_info else None,
        confidence  = round(lt_prob, 2),
    )

    structure = StructureResult(
        id         = struct_info[0]  if struct_info else None,
        name       = struct_info[1]  if struct_info else None,
        formula    = struct_info[13] if struct_info else None,
        cell_a     = struct_info[2]  if struct_info else None,
        cell_b     = struct_info[3]  if struct_info else None,
        cell_c     = struct_info[4]  if struct_info else None,
        sg_number  = struct_info[9]  if struct_info else None,
        sg_hm      = struct_info[11] if struct_info else None,
        confidence = round(st_prob, 2),
    )

    # ── Сохраняем GEOMETRIC результат в БД ───────────────────
    if session_id:
        try:
            _save_recognition_result(
                session_id      = session_id,
                lattice_type_id = lattice.id,
                structure_id    = structure.id,
                confidence      = lattice.confidence,
            )
        except Exception:
            pass

    ranking = sorted(
        [RankingItem(name_en=lt[2], name_ru=lt[1], prob=round(lt[3], 2)) for lt in lattice_names],
        key=lambda x: -x.prob,
    )

    return AnalyzeResponse(
        success         = True,
        session_id      = session_id,
        lattice         = lattice,
        structure       = structure,
        lattice_ranking = ranking,
        ml_results      = ml_results,
    )


@app.post("/api/chat", response_model=ChatResponse)
def chat(body: ChatRequest):
    """
    Чат с GigaChat.
    - Загружает богатый контекст структуры из БД по session_id
    - Принимает полную историю сообщений от фронта (обеспечивает continuity)
    - Сохраняет пару user/assistant в chat_message
    """
    if not _GC_READY:
        raise HTTPException(
            status_code=503,
            detail="GigaChat недоступен: проверьте GIGACHAT_CREDENTIALS в .env",
        )

    if not body.messages:
        raise HTTPException(status_code=400, detail="Список messages пуст")

    # ── Загружаем контекст из БД ──────────────────────────────
    try:
        ctx = _load_session_context(body.session_id)
    except Exception:
        ctx = {}

    system_prompt = _build_system_prompt(ctx)

    # ── Формируем payload для GigaChat ───────────────────────
    gc_messages = [Messages(role=MessagesRole.SYSTEM, content=system_prompt)]
    for m in body.messages:
        role = MessagesRole.USER if m.role == "user" else MessagesRole.ASSISTANT
        gc_messages.append(Messages(role=role, content=m.content))

    payload = Chat(
        model       = _GC_MODEL,
        messages    = gc_messages,
        max_tokens  = 1024,
        temperature = 0.5,
    )

    try:
        with GigaChat(
            credentials    = _GC_CREDENTIALS,
            scope          = _GC_SCOPE,
            verify_ssl_certs = False,
        ) as giga:
            response = giga.chat(payload)
            raw   = response.choices[0].message.content.strip()
            reply = _fix_gigachat_unicode(raw)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"GigaChat error: {e}")

    # ── Сохраняем в chat_message (fire-and-forget) ────────────
    user_msg = next((m.content for m in reversed(body.messages) if m.role == "user"), "")
    if user_msg:
        _save_chat_messages(body.session_id, user_msg, reply)

    return ChatResponse(reply=reply, model=_GC_MODEL)


@app.post("/api/export/docx")
def export_docx(body: AnalyzeResponse):
    """
    Генерирует минимальный DOCX-отчёт по результатам анализа.
    Принимает тот же объект, что возвращает /api/analyze.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx не установлен")

    doc = Document()

    # ── Заголовок ─────────────────────────────────────────────
    h = doc.add_heading("CRIS — Результат распознавания", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}").italic = True
    if body.session_id:
        doc.add_paragraph(f"Сессия: {body.session_id}")

    # ── По веществу ───────────────────────────────────────────
    doc.add_heading("По веществу", 1)

    p = doc.add_paragraph()
    p.add_run("База данных: ").bold = True
    if body.success and body.structure and body.structure.name:
        p.add_run(f"{body.structure.name} (уверенность: {body.structure.confidence:.2f})")
    else:
        p.add_run("совпадений не найдено")

    for mr in body.ml_results:
        if mr.category != "substance":
            continue
        labels = {"catboost_substance": "CatBoost · вещество", "rf": "Random Forest"}
        p = doc.add_paragraph()
        p.add_run(f"{labels.get(mr.method, mr.method)}: ").bold = True
        if mr.name_en:
            p.add_run(f"{mr.name_en} ({mr.confidence:.1f}%)")
            for item in mr.ranking[:3]:
                doc.add_paragraph(f"    {item.name_en}: {item.prob:.1f}%")
        else:
            p.add_run("нет данных")

    # ── Сингония ──────────────────────────────────────────────
    doc.add_heading("Сингония", 1)

    p = doc.add_paragraph()
    p.add_run("База данных: ").bold = True
    if body.success and body.lattice and body.lattice.name_en:
        p.add_run(f"{body.lattice.name_en} (уверенность: {body.lattice.confidence:.2f})")
    else:
        p.add_run("совпадений не найдено")

    for mr in body.ml_results:
        if mr.category != "crystal_system":
            continue
        p = doc.add_paragraph()
        p.add_run("CatBoost · сингония: ").bold = True
        if mr.name_en:
            p.add_run(f"{mr.name_en} ({mr.confidence:.1f}%)")
            for item in mr.ranking[:3]:
                doc.add_paragraph(f"    {item.name_en}: {item.prob:.1f}%")
        else:
            p.add_run("нет данных")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=cris_result.docx"},
    )
